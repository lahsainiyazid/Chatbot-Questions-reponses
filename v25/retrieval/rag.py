import json
import os
import io 
from PyPDF2 import PdfReader 
from docx import Document as DocxDocument
import certifi
import time
import hashlib
import redis
import pandas as pd
from fastapi import FastAPI,UploadFile,File,HTTPException
from pydantic import BaseModel
from docx import Document as DocxDocument

from pymongo import MongoClient
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_mongodb import MongoDBAtlasVectorSearch
from langchain_community.retrievers import BM25Retriever
from langchain_classic.retrievers import EnsembleRetriever
from sentence_transformers import CrossEncoder
from langchain_groq import ChatGroq
from dotenv import load_dotenv

# 1. Load environment variables from .env file
load_dotenv()

app = FastAPI()

# 2. REDIS CLOUD SETUP (Replaced local setup)
redis_url = os.getenv("UPSTASH_REDIS_URL")
if not redis_url:
    raise ValueError("REDIS_URL not found in .env file!")

redis_client = redis.from_url(
    redis_url,
    decode_responses=True,      # Keeps your json.loads/json.dumps logic working
    socket_connect_timeout=2,   # Kept your original timeouts
    socket_timeout=2
)
mongo_uri=os.getenv("MONGODB_URI")
mongo_client=MongoClient(mongo_uri,
                         tlsCAFile=certifi.where())
collection=mongo_client["rag_db"]["chunks"]
class QuestionRequest(BaseModel):
    question: str

llm_expansion = ChatGroq(model="llama-3.3-70b-instant", temperature=0, api_key=os.environ.get("GROQ_API_KEY"))
llm_answer = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, api_key=os.environ.get("GROQ_API_KEY"))

documents = []
try:
    cursor=collection.find({})
    for item in cursor:
        documents.append(Document(page_content=item.get("text",""),metadata={"source":item.get("source",""),"id":item.get("id","")}))
except Exception as e:
    print(f"Error:{e} while loading chunks!")
embeddings = HuggingFaceEmbeddings(
    model_name="intfloat/multilingual-e5-large",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True}
)

vector_store=MongoDBAtlasVectorSearch(collection=collection,embedding=embeddings,index_name="vector_index")
bm25 = BM25Retriever.from_documents(documents)
bm25.k = 5
dense = vector_store.as_retriever(search_kwargs={"k": 5})
hybrid = EnsembleRetriever(retrievers=[bm25, dense], weights=[0.6, 0.4])
reranker = CrossEncoder("BAAI/bge-reranker-base")
def sync_retrievers():
    global bm25,hybrid,dense 
    latest_docs=[]
    for item in collection.find({}):
        latest_docs.append(Document(page_content=item.get('text',''),metadata={'source':item.get('source',''),'id':item.get('id')}))
    if latest_docs:
        bm25=BM25Retriever.from_documents(latest_docs)
        bm25.k=5 
        dense=vector_store.as_retriever(search_kwargs={'k':5})
        hybrid=EnsembleRetriever(retrievers=[bm25,dense],weights=[0.6,0.4])
    else :
        bm25=None 
        dense=None
        hybrid=None 
def clear_cache():
    try:
        keys=redis_client.keys('rag_cache:*')
        if keys:
            redis_client.delete(*keys)
    except Exception as e:
        print(f"Error:{e} while clearing the cache!")

def parse_docx(content: bytes):
    try:
        doc = DocxDocument(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        chunks = []
        chunk_size = 1000
        words = text.split()
        for i in range(0, len(words), chunk_size):
            chunk_text = " ".join(words[i:i+chunk_size])
            if chunk_text.strip():
                chunks.append({
                    "text": chunk_text,
                    "id": hashlib.md5(chunk_text.encode()).hexdigest()
                })
        return chunks
    except Exception as e:
        print(f"Error {e} while loading docx")
        return []
def parse_pdf(content:bytes):
    try:
        pdf_reader=PdfReader(io.BytesIO(content))
        text=""
        for page in pdf_reader.pages:
            text+=page.extract_text()+'\n'
        chunks=[]
        chunk_size=1000
        words=text.split()
        for i in range(0,len(words),chunk_size):
            chunk_text=" ".join(words[i:i+chunk_size])
            if chunk_text.strip():
                chunks.append({"text":chunk_text,
                              "id":hashlib.md5(chunk_text.encode()).hexdigest()})
        return chunks 
    except Exception as e:
        print(f"Error:{e} while loading documents")
        return []
def parse_excel(content:bytes):
    try:
        excel_file=pd.ExcelFile(io.BytesIO(content))
        all_text=""
        for sheet_name in excel_file.sheet_names:
            df=pd.read_excel(excel_file,sheet_name=sheet_name)
            for col in df.columns:
                values=df[col].dropna().astype(str).to_list()
                if values:
                    all_text+=f"Column{col}"+",".join(values)+"\n"
        chunks=[]
        chunk_size=1000
        words=all_text.split()
        for i in range(0,len(words),chunk_size):
            chunk_text=" ".join(words[i:i+chunk_size])
            if chunk_text.strip():
                chunks.append({"text":chunk_text,"id":hashlib.md5(chunk_text.encode()).hexdigest()})
        return chunks 
    except Exception as e:
        print(f"Error {e} while loading the chunks")
        return []
@app.get("/")
def home():
    return {"Rag is running": "True", "Version": "V25"}
@app.get("/documents")
def get_documents():
    documents=list(collection.find({},{"_id":1,"source":1,"id":1,"text":1}))
    for doc in documents:
        doc["_id"]=str(doc["_id"]) #Because Json only supports standard types and "_id" is an object so we convert it first to string 
    return documents
@app.get("/document/{filename}")
def get_document(filename:str):
    docs=list(collection.find({"source":filename}))
    if not docs:
        raise HTTPException(404,detail="File not found")
    for doc in docs:
        doc["_id"]=str(doc["_id"])
    return {
        "filename":filename,
        "total_chunks":len(docs),
        "chunks":docs 

    }
    
@app.delete("/documents/{filename}")
def delete_document(filename:str):
    result=collection.delete_many({"source":filename})
    if result.deleted_count==0:
        raise HTTPException(404,detail="Document not found") #Code error for data not found 400 is for bad request 
    return{
        "message":"Document deleted successfully",
        "filename":filename,
        "deleted_chunks":result.deleted_count 
    }
@app.put("/document/{filename}")
async def update_document(filename:str,file:UploadFile=File(...)):
    if not (file.filename.endswith(".docx") or file.filename.endswith(".pdf") or file.filename.endswith(".xlsx") or file.filename.endswith(".xls")):
        raise HTTPException(400,detail="Only .docx,.pdf,.xls,.xlsx  files are accepted")
    delete_file=collection.delete_many({"source":filename})
    content=await file.read()
    if file.filename.endswith(".pdf"):
        chunks=parse_pdf(content)
    elif file.name.endswith(".docx") :
        chunks=parse_docx(content)
    else:
        chunks=parse_excel(content)
    if not chunks:
        raise HTTPException(400,detail="Error while extracting chunks from document!")
    for chunk in chunks:
        chunk["source"]=file.filename 
    result=collection.insert_many(chunks)
    return {
        "message":"Document updated successfully",
        "filename":file.filename,
        "new_chunks_count":len(result.inserted_ids),
        "deleted_chunks":delete_file.deleted_count,

    }

@app.post('/upload_file')
    async def upload_documents(file:UploadFile=File(...)):
    if not (file.filename.endswith('.docx') or file.filename.endswith('.pdf') or file.filename.endswith('.xls') or file.filename.endswith('.xlsx')):
       raise HTTPException(status_code=400,detail='Only .docx,.pdf,.xls,.xlsx files are supported')
    content=await file.read()
    if file.filename.endswith('.docx'):
        chunks=parse_docx(content)
    elif file.filename.endswith('.pdf'):
        chunks=parse_pdf(content)
    else:
        chunks=parse_excel(content)
    if not chunks:
        raise HTTPException(status_code=400,detail='Error while generating chunks')
    docs_to_add=[Document(page_content=chunk["text"],metadata={"source":chunk["source"],"id":chunk["id"]}) for chunk in chunks ]
    inserted_ids=vector_store.add_documents(docs_to_add)
    sync_retrievers()
    clear_cache()
    return {
    "Message":"Documents loaded successfully",
    "filename":file.filename,
    "Number of chunks":len(inserted_ids)
}
    

     
@app.post("/ask")
def ask_rag_question(request: QuestionRequest):
    start = time.time()
    normalized_question = request.question.strip().lower()
    cache_key = f"rag_cache:{hashlib.md5(normalized_question.encode('utf-8')).hexdigest()}"
    
    try:
        cached_response = redis_client.get(cache_key)
        if cached_response:
            print(f"{request.question} was found in the cache:")
            return json.loads(cached_response)
    except Exception as e:
        print(f"Error:{e}")

    expansion_start = time.time()
    expansion_prompt = f"""أنت محرك بحث ذكي متخصص في الوثائق الإدارية. مهمتك هي إعادة صياغة السؤال التالي إلى استعلامين بديلين دقيقين لتحسين البحث.
قواعد صارمة وممنوع مخالفتها:
أخرج الاستعلامين فقط، ولا شيء غيرهما.
ممنوع منعاً باتاً إضافة أي مقدمات أو خواتم أو شروحات.
ممنوع استخدام الأرقام أو النقاط (مثل: 1- أو *).
افصل بين الاستعلامين بسطر جديد واحد فقط.
استخدم مصطلحات إدارية رسمية ودقيقة.
السؤال: {request.question}
الاستعلامان:"""
    
    expanded_query = llm_expansion.invoke(expansion_prompt).content.strip()
    expansion_time = time.time() - expansion_start

    retrieval_start = time.time()
    results = hybrid.invoke(expanded_query)
    retrieval_time = time.time() - retrieval_start

    reranker_start = time.time()
    pairs = [(request.question, doc.page_content) for doc in results]
    scores = reranker.predict(pairs)
    ranked_docs = [doc for _, doc in sorted(zip(scores, results), key=lambda x: x[0], reverse=True)]
    final_docs = ranked_docs[:3]
    reranker_time = time.time() - reranker_start

    context = "\n\n".join(doc.page_content for doc in final_docs)
    system_prompt = """You are an expert assistant for Moroccan public administration.
Answer ONLY from the provided context.
Rules:
Never use external knowledge.
If the answer cannot be fully supported by the context, clearly say so.
Reply in the user's language.
"""
    user_prompt = f"""
Context:
{context}
Question:
{request.question}
"""

    llm_start = time.time()
    answer = llm_answer.invoke([("system", system_prompt), ("human", user_prompt)])
    llm_time = time.time() - llm_start
    total_time = time.time() - start

    token_usage = answer.response_metadata["token_usage"]
    final_response = {
        "Question": request.question,
        "Answer": answer.content.strip(),
        "total_time": round(total_time, 3),
        "expansion_time": round(expansion_time, 3),
        "retrieval_time": round(retrieval_time, 3),
        "reranker_time": round(reranker_time, 3),
        "llm_time": round(llm_time, 3),
        "token_usage": {
            "total_tokens": token_usage["total_tokens"],
            "prompt_tokens": token_usage["prompt_tokens"],
            "completion_tokens": token_usage["completion_tokens"]
        },
        "cached": "False"
    }

    try:
        redis_client.setex(cache_key, 86400, json.dumps(final_response))
    except Exception as e:
        print(f"Failed to save to cache error:{e}")

    return final_response
